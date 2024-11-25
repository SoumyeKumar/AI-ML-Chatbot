from docx import Document

# Create a new Document
doc = Document()

# Add a title to the document
doc.add_heading('Legal Motion Template', level=1)

# Add some instructional text for the user
doc.add_paragraph('This is a template for filling out a legal motion. Please replace the placeholders with your information.')

# Add placeholders for user input
doc.add_heading('1. Case Information', level=2)
doc.add_paragraph('Case Number: {{CaseNumber}}')
doc.add_paragraph('Case Name: {{CaseName}}')
doc.add_paragraph('Court Name: {{CourtName}}')

doc.add_heading('2. Parties Involved', level=2)
doc.add_paragraph('Plaintiff Name: {{PlaintiffName}}')
doc.add_paragraph('Defendant Name: {{DefendantName}}')

doc.add_heading('3. Motion Details', level=2)
doc.add_paragraph('Type of Motion: {{MotionType}}')
doc.add_paragraph('Motion Date: {{MotionDate}}')
doc.add_paragraph('Motion Description:\n{{MotionDescription}}')

doc.add_heading('4. Contact Information', level=2)
doc.add_paragraph('Attorney Name: {{AttorneyName}}')
doc.add_paragraph('Phone Number: {{PhoneNumber}}')
doc.add_paragraph('Email Address: {{EmailAddress}}')

# Add a footer for additional instructions
doc.add_paragraph('\nPlease ensure all placeholders are replaced before submitting this document.')

# Save the document locally
file_path = "Legal_Motion_Template_with_Placeholders.docx"
doc.save(file_path)
print(f"Template saved to {file_path}")
