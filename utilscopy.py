from vector_store import index
from model import embedding_model
from twilio.rest import Client
from datetime import datetime, timedelta
import threading
import googlemaps
import pandas as pd
import re
import docx
from io import BytesIO

# Twilio Configuration
TWILIO_ACCOUNT_SID = "ACa013e444a2b121bbaa4e474aacdade39"  # Replace with your Twilio Account SID
TWILIO_AUTH_TOKEN = "89d0372ab2e12caf3f3c368b6b6c3f9a"  # Replace with your Twilio Auth Token
TWILIO_PHONE_NUMBER = "+15079612120"  # Replace with your Twilio phone number

# Google Maps API Configuration
GOOGLE_MAPS_API_KEY = "AIzaSyC_35K1URATbFo0P0UBmb6-1706Gr3MIVA"  # Replace with your Google Maps API key
gmaps = googlemaps.Client(key=GOOGLE_MAPS_API_KEY)

# Initialize Twilio Client
twilio_client = Client(TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN)

# Load the legal cases dataset
LEGAL_CASES_FILE = "legal_cases_dataset.csv"  # Replace with the path to your dataset
legal_cases_df = pd.read_csv(LEGAL_CASES_FILE)

async def get_context_from_pinecone(query: str, top_k: int = 5) -> str:
    """
    Retrieves context from Pinecone for the given query.
    """
    query_embedding = embedding_model.embed_query(query)
    response = index.query(vector=query_embedding, top_k=top_k, include_metadata=True, namespace="legal-chatbot")
    context = " ".join([match['metadata']['output'] for match in response['matches']]) if response['matches'] else "No relevant context found."
    return context

def send_sms_reminder(to_phone, case_name, court_date):
    """
    Sends an SMS reminder using Twilio.
    """
    try:
        message = f"Reminder: You have a court date for '{case_name}' on {court_date}. Please prepare accordingly."
        twilio_client.messages.create(
            body=message,
            from_=TWILIO_PHONE_NUMBER,
            to=to_phone
        )
        print(f"SMS reminder sent to {to_phone}.")
    except Exception as e:
        print(f"Failed to send SMS: {e}")

def schedule_sms_reminder(to_phone, court_date, case_name):
    """
    Schedules an SMS reminder 24 hours before the court date.
    """
    try:
        court_datetime = datetime.strptime(court_date, "%Y-%m-%d %H:%M:%S")
        reminder_datetime = court_datetime - timedelta(days=1)
        current_datetime = datetime.now()
        time_until_reminder = (reminder_datetime - current_datetime).total_seconds()

        if time_until_reminder > 0:
            print(f"SMS reminder scheduled for {reminder_datetime}.")
            threading.Timer(time_until_reminder, send_sms_reminder, args=[to_phone, case_name, court_date]).start()
        else:
            print(f"Cannot schedule SMS reminder. The reminder time ({reminder_datetime}) is in the past.")
    except Exception as e:
        print(f"Error scheduling SMS reminder: {e}")

def geocode_location(location_name: str) -> str:
    """
    Converts a location name to coordinates using Google Maps Geocoding API.
    """
    try:
        geocode_result = gmaps.geocode(location_name)
        if geocode_result:
            lat = geocode_result[0]["geometry"]["location"]["lat"]
            lng = geocode_result[0]["geometry"]["location"]["lng"]
            return f"{lat},{lng}"
        else:
            print(f"No geocoding results for location: {location_name}")
            return ""
    except Exception as e:
        print(f"Error geocoding location: {e}")
        return ""

def find_nearby_lawyers(location: str, radius: int = 10000):
    """
    Finds nearby lawyers or legal services using Google Maps Places API.
    Default radius is 10,000 meters.
    """
    try:
        print(f"Searching for lawyers near location: {location} with radius: {radius}")
        results = gmaps.places_nearby(
            location=location,
            radius=radius,
            type="lawyer"  # Search for lawyers or legal services
        )
        places = []
        for place in results.get("results", []):
            name = place.get("name", "Unknown")
            address = place.get("vicinity", "Unknown address")
            url = f"https://www.google.com/maps/place/?q=place_id:{place['place_id']}"
            places.append({"name": name, "address": address, "url": url})
        return places
    except Exception as e:
        print(f"Error finding nearby lawyers: {e}")
        return []

def get_legal_resources():
    """
    Returns a list of trusted external legal resources.
    """
    return [
        {"name": "FindLaw", "url": "https://www.findlaw.com/"},
        {"name": "Legal Aid", "url": "https://www.lsc.gov/"},
        {"name": "American Bar Association", "url": "https://www.americanbar.org/"},
    ]

def lookup_case_by_number(case_number: str):
    """
    Looks up legal case details by case number in the dataset.
    """
    try:
        if legal_cases_df is None or legal_cases_df.empty:
            return "Error: Legal cases dataset is empty or not loaded properly."

        if "CaseNumber" not in legal_cases_df.columns:
            return "Error: The dataset does not contain a 'CaseNumber' column."

        case_details = legal_cases_df[legal_cases_df["CaseNumber"] == case_number]
        if not case_details.empty:
            return case_details.iloc[0].to_dict()
        else:
            return f"No case found with case number: {case_number}"
    except Exception as e:
        print(f"Error during case lookup: {e}")
        return f"An error occurred while looking up the case: {e}"

def schedule_appointment(to_phone, lawyer_name, appointment_time):
    """
    Schedules an appointment with a lawyer and sends SMS confirmation.
    """
    try:
        message = f"Appointment confirmed with {lawyer_name} on {appointment_time}. Please be on time."
        twilio_client.messages.create(
            body=message,
            from_=TWILIO_PHONE_NUMBER,
            to=to_phone
        )
        print(f"Appointment confirmation SMS sent to {to_phone}.")
    except Exception as e:
        print(f"Failed to send appointment confirmation SMS: {e}")

# Document Filling Assistance Functions
def extract_placeholders(text):
    """
    Extract placeholders in the format {{FieldName}} from a text.
    """
    return re.findall(r"\{\{(.*?)\}\}", text)

def read_docx(file):
    """
    Read and extract text from a DOCX file.
    """
    doc = docx.Document(file)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text

def fill_docx_template(file, user_inputs):
    """
    Fill placeholders in a DOCX file with user inputs.
    """
    doc = docx.Document(file)
    output_buffer = BytesIO()
    for paragraph in doc.paragraphs:
        for field, value in user_inputs.items():
            if f"{{{{{field}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{field}}}}}", value)

    # Save the filled document to buffer
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer
